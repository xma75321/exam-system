"""Word 文档文本提取器 — 使用 python-docx 读取段落和表格"""

from docx import Document


def extract_text(file_path: str) -> list[str]:
    """从 .docx 文件中提取所有非空文本行。

    Args:
        file_path: .docx 文件的绝对或相对路径

    Returns:
        非空文本行的列表（段落 + 表格单元格）
    """
    doc = Document(file_path)
    lines: list[str] = []

    # 提取段落文本
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    # 提取表格单元格文本
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    lines.append(text)

    return lines
