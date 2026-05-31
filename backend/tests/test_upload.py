"""文件上传功能测试"""

import io
import os
import pytest
from httpx import ASGITransport, AsyncClient

from app.main import app
from app.config import settings


@pytest.fixture
async def client():
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        yield ac


def create_test_docx() -> io.BytesIO:
    """使用 python-docx 创建一个有效的 .docx 测试文件"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("一、单选题（每题 2 分，共 20 分）")
    doc.add_paragraph("1. Python 是什么类型的语言？")
    doc.add_paragraph("A. 编译型")
    doc.add_paragraph("B. 解释型")
    doc.add_paragraph("答案：B")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


@pytest.mark.asyncio
async def test_upload_docx_success(client: AsyncClient):
    """上传 .docx 文件返回 200"""
    docx_file = create_test_docx()
    response = await client.post(
        "/api/upload",
        files={"file": ("test.docx", docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert response.status_code == 200
    data = response.json()
    assert data["code"] == 0
    assert data["data"]["filename"] == "test.docx"
    assert data["data"]["file_size"] > 0
    # 验证文件保存到磁盘
    assert os.path.exists(data["data"]["file_path"])


@pytest.mark.asyncio
async def test_upload_non_docx_file(client: AsyncClient):
    """上传非 .docx 文件返回 400"""
    response = await client.post(
        "/api/upload",
        files={"file": ("test.txt", b"not a docx", "text/plain")},
    )
    assert response.status_code == 400


@pytest.mark.asyncio
async def test_upload_no_file(client: AsyncClient):
    """无文件上传返回 422"""
    response = await client.post("/api/upload")
    assert response.status_code == 422


@pytest.mark.asyncio
async def test_upload_filename_is_uuid(client: AsyncClient):
    """验证保存的文件名包含 UUID"""
    docx_file = create_test_docx()
    response = await client.post(
        "/api/upload",
        files={"file": ("test.docx", docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    data = response.json()
    saved_name = os.path.basename(data["data"]["file_path"])
    # 应格式为 uuid.docx
    assert saved_name.endswith(".docx")
    assert len(saved_name) > 10  # UUID 长度远大于 10
